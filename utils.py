"""Utility functions for the Research Paper Reading Helper app.

The app stores everything in local files so it keeps working without a
database, login, or cloud service. DeepSeek summarisation is optional.
"""

from __future__ import annotations

import json
import os
import re
import urllib.error
import urllib.request
from collections import Counter
from datetime import date, datetime
from pathlib import Path
from typing import Any, Optional

import pandas as pd

try:
    from docx import Document
except ImportError:
    Document = None


DATA_DIR = Path(__file__).resolve().parent
SAVED_PAPERS_FILE = DATA_DIR / "saved_papers.json"
VOCABULARY_FILE = DATA_DIR / "vocabulary.csv"
EXPORTS_DIR = DATA_DIR / "exports"
DEEPSEEK_BASE_URL = "https://api.deepseek.com"
DEEPSEEK_MODEL = "deepseek-chat"
DEEPSEEK_MAX_CHARS = 60000

STOPWORDS = {
    "about",
    "after",
    "also",
    "among",
    "and",
    "are",
    "because",
    "been",
    "between",
    "both",
    "but",
    "can",
    "could",
    "did",
    "does",
    "during",
    "each",
    "for",
    "from",
    "had",
    "has",
    "have",
    "how",
    "however",
    "into",
    "may",
    "more",
    "most",
    "not",
    "our",
    "patients",
    "paper",
    "research",
    "result",
    "results",
    "show",
    "showed",
    "study",
    "such",
    "than",
    "that",
    "the",
    "their",
    "these",
    "this",
    "those",
    "through",
    "using",
    "was",
    "were",
    "which",
    "while",
    "with",
}

QUESTION_WORDS = {
    "a",
    "an",
    "any",
    "are",
    "based",
    "be",
    "by",
    "can",
    "did",
    "do",
    "does",
    "explain",
    "find",
    "for",
    "give",
    "how",
    "happen",
    "happened",
    "happens",
    "in",
    "is",
    "list",
    "me",
    "of",
    "on",
    "paper",
    "show",
    "study",
    "tell",
    "the",
    "there",
    "this",
    "to",
    "was",
    "were",
    "what",
    "when",
    "where",
    "whether",
    "which",
    "who",
    "why",
}

SYNONYM_GROUPS = {
    "aim": ["aim", "objective", "purpose", "goal"],
    "sample": ["sample", "participants", "subjects", "cells", "mice", "rats", "cohort", "population"],
    "method": ["method", "methods", "technique", "techniques", "procedure", "protocol", "assay", "experiment"],
    "measure": ["measure", "measured", "measurement", "assessed", "evaluated", "tested", "detected"],
    "result": ["result", "results", "finding", "findings", "outcome", "effect", "change"],
    "increase": ["increase", "increased", "higher", "elevated", "upregulated", "rise", "raised"],
    "decrease": ["decrease", "decreased", "lower", "reduced", "downregulated", "decline"],
    "significant": ["significant", "statistically", "p-value", "p<", "difference"],
    "association": ["association", "associated", "correlation", "correlated", "linked", "relationship"],
    "limitation": ["limitation", "limitations", "weakness", "bias", "constraint", "small sample", "limited"],
    "conclusion": ["conclusion", "concluded", "suggests", "indicates", "implies", "shows"],
    "source": ["source", "sources", "reference", "references", "citation", "bibliography"],
    "apoptosis": ["apoptosis", "programmed cell death", "cell death"],
    "inflammation": ["inflammation", "inflammatory", "immune response"],
    "biomarker": ["biomarker", "marker", "indicator"],
}

TERM_TO_SYNONYMS = {}
for synonym_group in SYNONYM_GROUPS.values():
    expanded_group = sorted(set(synonym_group))
    for synonym in expanded_group:
        TERM_TO_SYNONYMS[synonym] = expanded_group

SECTION_ALIASES = {
    "abstract": ["abstract", "summary"],
    "introduction": ["introduction", "background"],
    "methods": ["methods", "materials and methods", "methodology", "experimental procedures"],
    "results": ["results", "findings"],
    "discussion": ["discussion"],
    "conclusion": ["conclusion", "conclusions"],
    "limitations": ["limitations", "strengths and limitations"],
    "references": ["references", "bibliography", "literature cited", "works cited"],
}

VOCABULARY_COLUMNS = [
    "Term",
    "Simple explanation",
    "Related topic",
    "Example sentence",
    "Source paper title",
]

STARTER_TERMS = [
    {
        "Term": "Apoptosis",
        "Simple explanation": "programmed cell death",
        "Related topic": "Cell biology",
        "Example sentence": "Apoptosis removes damaged cells from the body.",
        "Source paper title": "Starter vocabulary",
    },
    {
        "Term": "Homeostasis",
        "Simple explanation": "maintenance of a stable internal environment",
        "Related topic": "Physiology",
        "Example sentence": "The body uses homeostasis to keep blood glucose stable.",
        "Source paper title": "Starter vocabulary",
    },
    {
        "Term": "Osmosis",
        "Simple explanation": "movement of water across a partially permeable membrane",
        "Related topic": "Cell transport",
        "Example sentence": "Water enters cells by osmosis.",
        "Source paper title": "Starter vocabulary",
    },
    {
        "Term": "Diffusion",
        "Simple explanation": "movement of particles from high concentration to low concentration",
        "Related topic": "Cell transport",
        "Example sentence": "Oxygen moves into cells by diffusion.",
        "Source paper title": "Starter vocabulary",
    },
    {
        "Term": "Enzyme",
        "Simple explanation": "protein that speeds up a chemical reaction",
        "Related topic": "Biochemistry",
        "Example sentence": "An enzyme helps break down the substrate faster.",
        "Source paper title": "Starter vocabulary",
    },
    {
        "Term": "ATP",
        "Simple explanation": "main energy molecule used by cells",
        "Related topic": "Cell metabolism",
        "Example sentence": "Muscle cells use ATP during contraction.",
        "Source paper title": "Starter vocabulary",
    },
    {
        "Term": "Mitosis",
        "Simple explanation": "cell division that produces two identical daughter cells",
        "Related topic": "Cell biology",
        "Example sentence": "Mitosis is important for growth and tissue repair.",
        "Source paper title": "Starter vocabulary",
    },
    {
        "Term": "Biomarker",
        "Simple explanation": "measurable sign of a biological condition",
        "Related topic": "Disease testing",
        "Example sentence": "Blood glucose can be used as a biomarker for diabetes.",
        "Source paper title": "Starter vocabulary",
    },
    {
        "Term": "Inflammation",
        "Simple explanation": "immune response to injury or infection",
        "Related topic": "Immunology",
        "Example sentence": "Inflammation can cause redness, heat, swelling, and pain.",
        "Source paper title": "Starter vocabulary",
    },
    {
        "Term": "Pathogen",
        "Simple explanation": "microorganism that can cause disease",
        "Related topic": "Microbiology",
        "Example sentence": "Some bacteria are pathogens that can infect humans.",
        "Source paper title": "Starter vocabulary",
    },
]


def ensure_storage_files() -> None:
    """Create the local storage files and export folder if they do not exist."""
    EXPORTS_DIR.mkdir(exist_ok=True)

    if not SAVED_PAPERS_FILE.exists():
        legacy_saved_file = Path.cwd() / "saved_papers.json"
        if legacy_saved_file != SAVED_PAPERS_FILE and legacy_saved_file.exists():
            SAVED_PAPERS_FILE.write_text(
                legacy_saved_file.read_text(encoding="utf-8"), encoding="utf-8"
            )
        else:
            SAVED_PAPERS_FILE.write_text("[]", encoding="utf-8")

    if not VOCABULARY_FILE.exists():
        pd.DataFrame(STARTER_TERMS, columns=VOCABULARY_COLUMNS).to_csv(
            VOCABULARY_FILE, index=False
        )


def load_saved_papers() -> list[dict[str, Any]]:
    """Load saved paper reviews from JSON, returning an empty list on errors."""
    ensure_storage_files()
    try:
        data = json.loads(SAVED_PAPERS_FILE.read_text(encoding="utf-8"))
        return data if isinstance(data, list) else []
    except (json.JSONDecodeError, OSError):
        return []


def save_paper(paper: dict[str, Any]) -> None:
    """Add or update one paper review in saved_papers.json."""
    ensure_storage_files()
    if not paper.get("paper_id"):
        paper["paper_id"] = make_paper_id(paper.get("paper_title", "untitled_paper"))
    papers = load_saved_papers()
    existing_index = next(
        (index for index, item in enumerate(papers) if item.get("paper_id") == paper.get("paper_id")),
        None,
    )

    if existing_index is None:
        papers.append(paper)
    else:
        papers[existing_index] = paper

    SAVED_PAPERS_FILE.write_text(
        json.dumps(papers, indent=2, ensure_ascii=False), encoding="utf-8"
    )


def delete_paper(paper_id: str) -> None:
    """Delete a saved paper review by its paper ID."""
    papers = [paper for paper in load_saved_papers() if paper.get("paper_id") != paper_id]
    SAVED_PAPERS_FILE.write_text(
        json.dumps(papers, indent=2, ensure_ascii=False), encoding="utf-8"
    )


def load_vocabulary() -> pd.DataFrame:
    """Load the vocabulary CSV and repair missing columns if needed."""
    ensure_storage_files()
    try:
        vocabulary = pd.read_csv(VOCABULARY_FILE).fillna("")
    except (pd.errors.EmptyDataError, OSError):
        vocabulary = pd.DataFrame(columns=VOCABULARY_COLUMNS)

    for column in VOCABULARY_COLUMNS:
        if column not in vocabulary.columns:
            vocabulary[column] = ""
    return vocabulary[VOCABULARY_COLUMNS]


def save_vocabulary(vocabulary: pd.DataFrame) -> None:
    """Save vocabulary terms to vocabulary.csv."""
    ensure_storage_files()
    vocabulary[VOCABULARY_COLUMNS].to_csv(VOCABULARY_FILE, index=False)


def clean_filename(text: str) -> str:
    """Turn a paper title into a safe file-name section."""
    cleaned = re.sub(r"[^A-Za-z0-9_-]+", "_", text.strip().lower())
    cleaned = re.sub(r"_+", "_", cleaned).strip("_")
    return cleaned[:80] or "untitled_paper"


def make_paper_id(title: str) -> str:
    """Create a readable unique paper ID using title plus the current time."""
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{clean_filename(title)}_{stamp}"


def serialise_value(value: Any) -> Any:
    """Convert dates and other non-JSON values into JSON-friendly values."""
    if isinstance(value, (date, datetime)):
        return value.isoformat()
    return value


def extract_text_from_upload(uploaded_file: Any) -> str:
    """Read text from an uploaded TXT or PDF file.

    PDF support uses pypdf. If pypdf is not installed, the app returns a clear
    message instead of crashing.
    """
    if uploaded_file is None:
        return ""

    file_name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()
    if file_name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")

    if file_name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
        except ImportError:
            return "PDF reading requires pypdf. Install it with: pip install pypdf"

        try:
            import io

            reader = PdfReader(io.BytesIO(data))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages).strip()
        except Exception as error:
            return f"Could not read this PDF: {error}"

    return data.decode("utf-8", errors="ignore")


def clean_paper_text(text: str) -> str:
    """Normalise whitespace from pasted or uploaded paper text."""
    text = text.replace("\r", "\n")
    text = re.sub(r"-\n", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def split_sentences(text: str) -> list[str]:
    """Split paper text into readable sentence-like chunks."""
    cleaned = re.sub(r"\s+", " ", text.strip())
    if not cleaned:
        return []
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9])", cleaned)
    return [sentence.strip() for sentence in sentences if len(sentence.strip()) > 25]


def extract_sections(text: str) -> dict[str, str]:
    """Extract common paper sections using simple local heading detection."""
    cleaned = clean_paper_text(text)
    sections = {key: "" for key in SECTION_ALIASES}
    if not cleaned:
        return sections

    heading_pattern = re.compile(
        r"(?im)^\s*(abstract|summary|introduction|background|methods|materials and methods|"
        r"methodology|experimental procedures|results|findings|discussion|conclusion|"
        r"conclusions|limitations|strengths and limitations|references|bibliography|"
        r"literature cited|works cited)\s*$"
    )
    matches = list(heading_pattern.finditer(cleaned))

    if not matches:
        sentences = split_sentences(cleaned)
        sections["abstract"] = " ".join(sentences[:8])
        return sections

    for index, match in enumerate(matches):
        raw_heading = match.group(1).lower()
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(cleaned)
        content = cleaned[start:end].strip()
        for section, aliases in SECTION_ALIASES.items():
            if raw_heading in aliases:
                sections[section] = content
                break
    return sections


def extract_sources(text: str) -> list[str]:
    """Extract likely source/reference entries from a paper's reference list."""
    sections = extract_sections(text)
    references = sections.get("references", "")
    if not references:
        match = re.search(
            r"(?is)\n\s*(references|bibliography|literature cited|works cited)\s*\n(.+)$",
            clean_paper_text(text),
        )
        references = match.group(2).strip() if match else ""

    if not references:
        return []

    lines = [line.strip() for line in references.splitlines() if line.strip()]
    sources = []
    current = ""
    new_entry = re.compile(r"^(\[\d+\]|\d+[\.)]\s+|[A-Z][A-Za-z'-]+,\s+[A-Z])")

    for line in lines:
        lowered = line.lower()
        if lowered in {"appendix", "appendices", "supplementary material"}:
            break
        if new_entry.match(line) and current:
            sources.append(current.strip())
            current = line
        else:
            current = f"{current} {line}".strip()

    if current:
        sources.append(current.strip())

    cleaned_sources = []
    for source in sources:
        source = re.sub(r"\s+", " ", source).strip()
        if len(source) > 20 and source.lower() not in {"references", "bibliography"}:
            cleaned_sources.append(source)
    return cleaned_sources


def format_citation_source(source: str, style: str, index: int) -> str:
    """Format a plain extracted reference in a best-effort citation style."""
    cleaned = re.sub(r"\s+", " ", source).strip()
    cleaned = re.sub(r"^(\[\d+\]|\d+[\.)])\s*", "", cleaned).strip()
    cleaned = cleaned.rstrip(".")
    style = style.lower()

    if style == "original":
        return source.strip()
    if style == "vancouver":
        return f"{index}. {cleaned}."
    if style == "mla":
        return f"{cleaned}."
    if style == "harvard":
        return f"{cleaned}."
    if style == "apa":
        return f"{cleaned}."
    return source.strip()


def format_citation_sources(sources: list[str], style: str) -> list[str]:
    """Format extracted sources for display or download."""
    return [
        format_citation_source(source, style, index)
        for index, source in enumerate(sources, start=1)
        if str(source).strip()
    ]


def _infer_citation_author(source: str) -> str:
    """Infer the first author surname from a plain reference string."""
    cleaned = re.sub(r"^(\[\d+\]|\d+[\.)])\s*", "", source).strip()
    if not cleaned:
        return "Author"
    first_chunk = re.split(r"\.|,", cleaned, maxsplit=1)[0].strip()
    words = re.findall(r"[A-Z][A-Za-z'-]+", first_chunk)
    return words[0] if words else "Author"


def _infer_citation_year(source: str) -> str:
    """Infer a publication year from a plain reference string."""
    match = re.search(r"\b(19|20)\d{2}\b", source)
    return match.group(0) if match else "n.d."


def make_in_text_citation(source: str, style: str, index: int) -> str:
    """Create a best-effort in-text citation for one source."""
    style = style.lower()
    author = _infer_citation_author(source)
    year = _infer_citation_year(source)

    if style == "vancouver":
        return f"[{index}]"
    if style == "mla":
        return f"({author})"
    if style in {"apa", "harvard", "original"}:
        return f"({author}, {year})"
    return f"({author}, {year})"


def extract_keywords(text: str, limit: int = 12) -> list[str]:
    """Find likely important biomedical keywords using local word frequency."""
    words = re.findall(r"\b[A-Za-z][A-Za-z-]{3,}\b", text.lower())
    filtered = [
        word.strip("-")
        for word in words
        if word not in STOPWORDS and not word.endswith("ing") and len(word) > 3
    ]
    counts = Counter(filtered)
    return [word for word, _ in counts.most_common(limit)]


def summarise_text(text: str, sentence_count: int = 4) -> str:
    """Create a short extractive summary from the most informative sentences."""
    sentences = split_sentences(text)
    if not sentences:
        return ""

    keywords = set(extract_keywords(text, limit=20))
    scored = []
    for position, sentence in enumerate(sentences):
        words = re.findall(r"\b[A-Za-z][A-Za-z-]{3,}\b", sentence.lower())
        keyword_hits = sum(1 for word in words if word in keywords)
        result_hits = len(
            re.findall(
                r"\b(increased|decreased|significant|associated|correlated|reduced|higher|lower|improved)\b",
                sentence.lower(),
            )
        )
        score = keyword_hits + (2 * result_hits) - (position * 0.02)
        scored.append((score, position, sentence))

    best = sorted(scored, reverse=True)[:sentence_count]
    ordered = sorted(best, key=lambda item: item[1])
    return " ".join(sentence for _, _, sentence in ordered)


def simplify_sentence(sentence: str, max_words: int = 22) -> str:
    """Shorten and lightly rephrase a sentence for beginner summaries."""
    sentence = re.sub(r"\([^)]*\)", "", sentence)
    sentence = re.sub(r"\s+", " ", sentence).strip(" .;:")
    replacements = {
        "demonstrated": "showed",
        "indicated": "suggested",
        "utilized": "used",
        "significantly": "clearly",
        "associated with": "linked to",
        "correlated with": "linked to",
        "investigated": "studied",
        "evaluate": "test",
        "evaluated": "tested",
        "assessed": "measured",
    }
    lowered = sentence.lower()
    for original, replacement in replacements.items():
        lowered = lowered.replace(original, replacement)
    words = lowered.split()
    if len(words) > max_words:
        lowered = " ".join(words[:max_words]).rstrip(",;:") + "..."
    return lowered[:1].upper() + lowered[1:] if lowered else ""


def strip_template_prefix(text: str) -> str:
    """Remove repeated template prefixes from bullet points."""
    prefixes = [
        "This paper focuses on ",
        "The researchers used ",
        "The main finding was that ",
        "Overall, the study suggests that ",
        "A key caution is that ",
    ]
    for prefix in prefixes:
        if text.startswith(prefix):
            text = text[len(prefix) :]
            break
    return text[:1].upper() + text[1:] if text else text


def first_relevant_sentence(text: str, terms: list[str]) -> str:
    """Return the first sentence containing any target term."""
    for sentence in split_sentences(text):
        lowered = sentence.lower()
        if any(term in lowered for term in terms):
            return sentence
    sentences = split_sentences(text)
    return sentences[0] if sentences else ""


def generate_reworded_summary(sections: dict[str, str], full_text: str) -> str:
    """Generate a short reworded summary using local templates."""
    abstract_or_intro = sections.get("abstract") or sections.get("introduction") or full_text
    methods_text = sections.get("methods") or full_text
    results_text = sections.get("results") or sections.get("abstract") or full_text
    discussion_text = sections.get("discussion") or sections.get("conclusion") or full_text

    focus = simplify_sentence(
        first_relevant_sentence(
            abstract_or_intro,
            ["aim", "objective", "purpose", "investigat", "study", "test", "whether"],
        )
    )
    method = simplify_sentence(
        first_relevant_sentence(
            methods_text,
            ["method", "sample", "participant", "cell", "mice", "assay", "measured", "analysis", "trial"],
        )
    )
    finding = simplify_sentence(
        first_relevant_sentence(
            results_text,
            ["increased", "decreased", "significant", "higher", "lower", "associated", "correlated", "found", "showed"],
        )
    )
    meaning = simplify_sentence(
        first_relevant_sentence(
            discussion_text,
            ["suggest", "indicat", "conclude", "important", "therefore", "may"],
        )
    )
    limitation = simplify_sentence(
        first_relevant_sentence(
            full_text,
            ["limitation", "limited", "small sample", "bias", "future", "further research"],
        )
    )

    parts = []
    if focus:
        parts.append(f"This paper focuses on {focus[0].lower() + focus[1:]}")
    if finding:
        parts.append(f"The main finding was that {finding[0].lower() + finding[1:]}")
    if limitation:
        parts.append(f"A key caution is that {limitation[0].lower() + limitation[1:]}")
    if not limitation and method:
        parts.append(f"The researchers used {method[0].lower() + method[1:]}")
    if len(parts) < 3 and meaning:
        parts.append(f"Overall, the study suggests that {meaning[0].lower() + meaning[1:]}")

    if not parts:
        return summarise_text(full_text, sentence_count=1)
    return " ".join(part.rstrip(".") + "." for part in parts[:3])


def generate_reworded_points(sections: dict[str, str], full_text: str) -> list[str]:
    """Create beginner-friendly summary bullets without copying whole passages."""
    points = []
    summary = generate_reworded_summary(sections, full_text)
    for sentence in split_sentences(summary):
        points.append(strip_template_prefix(sentence))
    if len(points) < 4:
        keywords = extract_keywords(full_text, limit=6)
        if keywords:
            points.append(f"Important recurring topics include {', '.join(keywords[:5])}.")
    return points[:4]


def extract_key_points(text: str, limit: int = 8) -> list[str]:
    """Return bullet-style reading points from the paper text."""
    summary = summarise_text(text, sentence_count=limit)
    return split_sentences(summary)[:limit]


def _clean_deepseek_list(value: Any, limit: int = 8) -> list[str]:
    """Normalise DeepSeek JSON list fields into display-safe strings."""
    if isinstance(value, str):
        items = [line.strip(" -*\t") for line in value.splitlines()]
    elif isinstance(value, list):
        items = [str(item).strip(" -*\t") for item in value]
    else:
        return []
    return [item for item in items if item][:limit]


def _deepseek_endpoint() -> str:
    """Return the configured DeepSeek chat completions endpoint."""
    base_url = os.environ.get("DEEPSEEK_BASE_URL", DEEPSEEK_BASE_URL).rstrip("/")
    if base_url.endswith("/chat/completions"):
        return base_url
    return f"{base_url}/chat/completions"


def _deepseek_chat_json(
    prompt: str,
    system_message: str,
    max_tokens: int = 1400,
) -> tuple[dict[str, Any], str] | None:
    """Call DeepSeek's OpenAI-compatible chat API and parse a JSON response."""
    api_key = os.environ.get("DEEPSEEK_API_KEY", "").strip()
    if not api_key:
        return None

    model = os.environ.get("DEEPSEEK_MODEL", DEEPSEEK_MODEL).strip() or DEEPSEEK_MODEL
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_message},
            {"role": "user", "content": prompt},
        ],
        "response_format": {"type": "json_object"},
        "temperature": 0.2,
        "max_tokens": max_tokens,
        "stream": False,
    }
    request = urllib.request.Request(
        _deepseek_endpoint(),
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )

    try:
        with urllib.request.urlopen(request, timeout=45) as response:
            raw_response = response.read().decode("utf-8")
        response_data = json.loads(raw_response)
        content = response_data["choices"][0]["message"]["content"]
        generated = json.loads(content)
    except (KeyError, IndexError, json.JSONDecodeError, OSError, urllib.error.URLError):
        return None

    return generated, model


def call_deepseek_summary(cleaned_text: str) -> dict[str, Any] | None:
    """Ask DeepSeek for a structured beginner-friendly paper summary.

    The app remains local-first: this function only runs when DEEPSEEK_API_KEY
    is present, and any API or parsing problem falls back to the local summary.
    """
    try:
        max_chars = int(os.environ.get("DEEPSEEK_MAX_CHARS", DEEPSEEK_MAX_CHARS))
    except ValueError:
        max_chars = DEEPSEEK_MAX_CHARS
    paper_excerpt = cleaned_text[:max(4000, max_chars)]

    prompt = f"""
Read the research paper text below and return only valid JSON with these keys:
short_summary: one short beginner-friendly paragraph that explains the research question, approach, main finding, and why it matters.
key_points: 5 concise content points covering the aim, methods, main results, interpretation, and study caution.
keywords: 8 important biomedical or research terms.
method_points: up to 5 points about methods, samples, assays, measures, or analysis.
result_points: up to 5 points about findings or results.
limitation_points: up to 5 limitations or cautions stated or strongly implied by the paper.

Do not invent citations, statistics, or claims. Use simple wording suitable for
a first-year biomedical science student.

Paper text:
{paper_excerpt}
""".strip()

    result = _deepseek_chat_json(
        prompt,
        "You summarise research papers accurately and return strict JSON.",
        max_tokens=1600,
    )
    if not result:
        return None

    generated, model = result
    short_summary = str(generated.get("short_summary", "")).strip()
    if not short_summary:
        return None

    return {
        "short_summary": short_summary,
        "key_points": _clean_deepseek_list(generated.get("key_points"), limit=5),
        "keywords": _clean_deepseek_list(generated.get("keywords"), limit=8),
        "method_points": _clean_deepseek_list(generated.get("method_points"), limit=5),
        "result_points": _clean_deepseek_list(generated.get("result_points"), limit=5),
        "limitation_points": _clean_deepseek_list(generated.get("limitation_points"), limit=5),
        "summary_provider": f"DeepSeek ({model})",
    }


def call_deepseek_comparison(papers: list[dict[str, Any]]) -> dict[str, Any] | None:
    """Ask DeepSeek to compare the selected papers using their extracted content."""
    if len(papers) < 2:
        return None

    paper_briefs = []
    for index, paper in enumerate(papers, start=1):
        title = paper.get("paper_title") or paper.get("uploaded_file_name") or f"Paper {index}"
        brief = {
            "title": title,
            "summary": paper.get("auto_summary", ""),
            "key_points": paper.get("auto_key_points", []),
            "keywords": paper.get("auto_keywords", []),
            "methods": paper.get("auto_method_points", []),
            "results": paper.get("auto_result_points", []),
            "limitations": paper.get("auto_limitation_points", []),
            "text_excerpt": clean_paper_text(paper.get("paper_text", ""))[:8000],
        }
        paper_briefs.append(brief)

    prompt = f"""
Compare these research papers and return only valid JSON with these keys:
comparison_summary: one concise paragraph explaining the most important similarities and differences.
shared_themes: 3 to 6 shared themes across the papers.
key_differences: 3 to 6 important differences between papers.
paper_takeaways: an object where each key is a paper title and each value is a list of 2 to 4 key takeaways.
possible_oppositions: up to 4 places where findings, populations, methods, or interpretations may differ or conflict.
study_cautions: up to 5 cautions about comparing these papers fairly.

Use only the supplied paper content. Do not invent citations, statistics, sample
sizes, or findings. Keep the wording clear for a first-year biomedical science
student.

Papers:
{json.dumps(paper_briefs, ensure_ascii=False)}
""".strip()

    result = _deepseek_chat_json(
        prompt,
        "You compare research papers accurately and return strict JSON.",
        max_tokens=1800,
    )
    if not result:
        return None

    generated, model = result
    comparison_summary = str(generated.get("comparison_summary", "")).strip()
    if not comparison_summary:
        return None

    raw_takeaways = generated.get("paper_takeaways", {})
    paper_takeaways = {}
    if isinstance(raw_takeaways, dict):
        for title, takeaways in raw_takeaways.items():
            paper_takeaways[str(title)] = _clean_deepseek_list(takeaways, limit=4)

    return {
        "comparison_summary": comparison_summary,
        "shared_themes": _clean_deepseek_list(generated.get("shared_themes"), limit=6),
        "key_differences": _clean_deepseek_list(generated.get("key_differences"), limit=6),
        "paper_takeaways": paper_takeaways,
        "possible_oppositions": _clean_deepseek_list(generated.get("possible_oppositions"), limit=4),
        "study_cautions": _clean_deepseek_list(generated.get("study_cautions"), limit=5),
        "comparison_provider": f"DeepSeek ({model})",
    }


def find_sentences_with_terms(text: str, terms: list[str], limit: int = 5) -> list[str]:
    """Find sentences that mention any of the provided terms."""
    sentences = split_sentences(text)
    found = []
    for sentence in sentences:
        lowered = sentence.lower()
        if any(term.lower() in lowered for term in terms):
            found.append(sentence)
        if len(found) >= limit:
            break
    return found


def recall_keyword_matches(paper: dict[str, Any], query: str, limit: int = 20) -> list[dict[str, str]]:
    """Search paper content, extracted points, and sources for keyword matches."""
    keywords = [word.lower() for word in re.findall(r"\b[\w-]+\b", query) if len(word) > 1]
    if not keywords:
        return []

    search_blocks = [
        ("Paper text", paper.get("paper_text", "")),
        ("Short summary", paper.get("auto_summary", "")),
        ("Content points", "\n".join(paper.get("auto_key_points", []))),
        ("Methods points", "\n".join(paper.get("auto_method_points", []))),
        ("Results points", "\n".join(paper.get("auto_result_points", []))),
        ("Limitations points", "\n".join(paper.get("auto_limitation_points", []))),
        ("Sources", "\n".join(paper.get("sources", []))),
    ]

    matches = []
    seen = set()
    for section, text in search_blocks:
        if not text:
            continue
        chunks = split_sentences(text)
        if not chunks:
            chunks = [line.strip() for line in text.splitlines() if line.strip()]
        for chunk in chunks:
            lowered = chunk.lower()
            matched_terms = [keyword for keyword in keywords if keyword in lowered]
            if not matched_terms:
                continue
            key = (section, chunk)
            if key in seen:
                continue
            seen.add(key)
            matches.append(
                {
                    "section": section,
                    "matched_terms": ", ".join(sorted(set(matched_terms))),
                    "snippet": chunk,
                }
            )
            if len(matches) >= limit:
                return matches
    return matches


def recall_answer_question(paper: dict[str, Any], question: str, limit: int = 5) -> dict[str, Any]:
    """Answer a natural-language question using only local paper text.

    This is extractive rather than generative: the answer is assembled from the
    most relevant sentences found in the paper and extracted sections.
    """
    query_terms = [
        word.lower()
        for word in re.findall(r"\b[A-Za-z][A-Za-z-]{2,}\b", question)
        if word.lower() not in QUESTION_WORDS and word.lower() not in STOPWORDS
    ]
    expanded_terms = sorted(
        {
            related_term
            for term in query_terms
            for related_term in TERM_TO_SYNONYMS.get(term, [term])
        }
    )
    if not query_terms:
        return {
            "answer": "Enter a more specific question or keyword phrase.",
            "main_point": "",
            "answer_groups": {},
            "key_details": [],
            "confidence": "Low",
            "matches": [],
        }

    search_blocks = [
        ("Short summary", paper.get("auto_summary", "")),
        ("Content points", "\n".join(paper.get("auto_key_points", []))),
        ("Methods points", "\n".join(paper.get("auto_method_points", []))),
        ("Results points", "\n".join(paper.get("auto_result_points", []))),
        ("Limitations points", "\n".join(paper.get("auto_limitation_points", []))),
        ("Paper text", paper.get("paper_text", "")),
        ("Sources", "\n".join(paper.get("sources", []))),
    ]

    scored_matches = []
    seen = set()
    for section, text in search_blocks:
        if not text:
            continue
        chunks = split_sentences(text)
        if not chunks:
            chunks = [line.strip() for line in text.splitlines() if line.strip()]
        for chunk in chunks:
            lowered = chunk.lower()
            matched_terms = [term for term in expanded_terms if term in lowered]
            if not matched_terms:
                continue
            phrase_bonus = 2 if " ".join(query_terms) in lowered else 0
            section_bonus = 1 if section != "Paper text" else 0
            exact_bonus = sum(1 for term in query_terms if term in lowered)
            score = len(set(matched_terms)) + exact_bonus + phrase_bonus + section_bonus
            key = (section, chunk)
            if key in seen:
                continue
            seen.add(key)
            scored_matches.append(
                {
                    "score": score,
                    "section": section,
                    "matched_terms": ", ".join(sorted(set(matched_terms))),
                    "snippet": chunk,
                }
            )

    if not scored_matches:
        return {
            "answer": "I could not find an answer in the uploaded paper text.",
            "main_point": "",
            "answer_groups": {},
            "key_details": [],
            "confidence": "Low",
            "matches": [],
        }

    ranked = sorted(scored_matches, key=lambda item: item["score"], reverse=True)[:limit]
    concise_points = []
    for match in ranked:
        point = re.sub(r"^\s*[-*•]\s*", "", match["snippet"]).strip()
        point = re.sub(r"\s+", " ", point)
        if len(point) > 170:
            point = point[:167].rsplit(" ", 1)[0] + "..."
        if point and point not in concise_points:
            concise_points.append(point)

    max_score = ranked[0]["score"]
    confidence = "High" if max_score >= 4 else "Medium" if max_score >= 2 else "Low"
    terms_text = ", ".join(sorted(set(query_terms[:4])))
    main_point = concise_points[0] if concise_points else ""
    answer_groups = {}
    if main_point:
        answer_groups["Main answer"] = [main_point]
    if len(concise_points) > 1:
        answer_groups["Key details"] = concise_points[1:4]

    return {
        "answer": f"Answer about {terms_text}.",
        "main_point": main_point,
        "answer_groups": answer_groups,
        "key_details": answer_groups.get("Key details", []),
        "confidence": confidence,
        "matches": [
            {
                "score": match["score"],
                "section": match["section"],
                "matched_terms": match["matched_terms"],
                "snippet": match["snippet"],
            }
            for match in ranked
        ],
    }


def recall_answer_across_papers(
    papers: list[dict[str, Any]],
    question: str,
    limit: int = 8,
) -> dict[str, Any]:
    """Answer a recall question using multiple uploaded papers."""
    usable_papers = [
        paper
        for paper in papers
        if paper.get("paper_text", "").strip() or paper.get("auto_summary", "").strip()
    ]
    if not usable_papers:
        return {
            "answer": "Upload or paste at least one paper before using recall.",
            "main_point": "",
            "answer_groups": {},
            "key_details": [],
            "confidence": "Low",
            "matches": [],
        }

    combined_matches = []
    confidence_scores = {"Low": 1, "Medium": 2, "High": 3}
    best_confidence_score = 1
    for index, paper in enumerate(usable_papers, start=1):
        title = paper.get("paper_title") or paper.get("uploaded_file_name") or f"Paper {index}"
        result = recall_answer_question(paper, question, limit=limit)
        best_confidence_score = max(
            best_confidence_score,
            confidence_scores.get(result.get("confidence", "Low"), 1),
        )
        for match in result.get("matches", []):
            combined_matches.append(
                {
                    **match,
                    "paper_title": title,
                    "score": match.get("score", 0),
                }
            )

    if not combined_matches:
        return {
            "answer": "I could not find an answer in the selected papers.",
            "main_point": "",
            "answer_groups": {},
            "key_details": [],
            "confidence": "Low",
            "matches": [],
        }

    ranked = sorted(combined_matches, key=lambda item: item.get("score", 0), reverse=True)[:limit]
    concise_points = []
    for match in ranked:
        point = re.sub(r"^\s*[-*•]\s*", "", match["snippet"]).strip()
        point = re.sub(r"\s+", " ", point)
        if len(point) > 170:
            point = point[:167].rsplit(" ", 1)[0] + "..."
        labelled_point = f"{match['paper_title']}: {point}"
        if labelled_point not in concise_points:
            concise_points.append(labelled_point)

    query_terms = [
        word.lower()
        for word in re.findall(r"\b[A-Za-z][A-Za-z-]{2,}\b", question)
        if word.lower() not in QUESTION_WORDS and word.lower() not in STOPWORDS
    ]
    terms_text = ", ".join(sorted(set(query_terms[:4])))
    confidence = {1: "Low", 2: "Medium", 3: "High"}.get(best_confidence_score, "Low")

    return {
        "answer": f"Answer about {terms_text}.",
        "main_point": concise_points[0] if concise_points else "",
        "answer_groups": {"Main answer": concise_points[:1], "Key details": concise_points[1:4]},
        "key_details": concise_points[1:4],
        "confidence": confidence,
        "matches": [
            {
                "paper_title": match["paper_title"],
                "section": match["section"],
                "matched_terms": match["matched_terms"],
                "snippet": match["snippet"],
            }
            for match in ranked
        ],
    }


def compare_papers(papers: list[dict[str, Any]]) -> dict[str, Any]:
    """Compare multiple papers using extracted local summaries and keywords."""
    def meaningful_terms(text: str) -> set[str]:
        return {
            word
            for word in re.findall(r"\b[A-Za-z][A-Za-z-]{3,}\b", text.lower())
            if word not in STOPWORDS and word not in QUESTION_WORDS
        }

    def theme_phrase_from_terms(terms: set[str]) -> str:
        priority_terms = [
            term
            for term in sorted(terms)
            if term in TERM_TO_SYNONYMS
            or term in {"cells", "treatment", "apoptosis", "biomarker", "inflammation", "sample", "method", "result"}
        ]
        chosen = priority_terms[:3] or sorted(terms)[:3]
        return " / ".join(chosen)

    prepared = []
    for paper in papers:
        title = paper.get("paper_title") or paper.get("uploaded_file_name") or "Untitled paper"
        keywords = set(paper.get("auto_keywords", []))
        text_for_keywords = " ".join(
            [
                paper.get("auto_summary", ""),
                " ".join(paper.get("auto_key_points", [])),
                " ".join(paper.get("auto_result_points", [])),
            ]
        )
        if not keywords and text_for_keywords.strip():
            keywords = set(extract_keywords(text_for_keywords, limit=12))
        prepared.append(
            {
                "title": title,
                "summary": paper.get("auto_summary", ""),
                "keywords": keywords,
                "ideas": [
                    sentence
                    for sentence in split_sentences(
                        " ".join(
                            [
                                paper.get("auto_summary", ""),
                                " ".join(paper.get("auto_key_points", [])),
                                " ".join(paper.get("auto_result_points", [])),
                                " ".join(paper.get("auto_limitation_points", [])),
                            ]
                        )
                    )
                ][:8],
                "methods": paper.get("auto_method_points", []),
                "results": paper.get("auto_result_points", []),
                "limitations": paper.get("auto_limitation_points", []),
            }
        )

    if not prepared:
        return {
            "papers": [],
            "shared_keywords": [],
            "unique_keywords": {},
            "oppositions": [],
            "comparison_provider": "Local extractor",
        }

    keyword_sets = [item["keywords"] for item in prepared if item["keywords"]]
    shared_keywords = sorted(set.intersection(*keyword_sets)) if keyword_sets else []
    all_keywords = set.union(*keyword_sets) if keyword_sets else set()
    unique_keywords = {}
    for item in prepared:
        other_keywords = set()
        for other in prepared:
            if other is not item:
                other_keywords.update(other["keywords"])
        unique_keywords[item["title"]] = sorted(item["keywords"] - other_keywords)

    common_ideas = []
    for first_index, first in enumerate(prepared):
        for second in prepared[first_index + 1 :]:
            for first_idea in first["ideas"]:
                first_terms = meaningful_terms(first_idea)
                if not first_terms:
                    continue
                for second_idea in second["ideas"]:
                    second_terms = meaningful_terms(second_idea)
                    shared_terms = first_terms & second_terms
                    if len(shared_terms) < 2:
                        continue
                    common_ideas.append(
                        {
                            "theme": theme_phrase_from_terms(shared_terms),
                            "papers": f"{first['title']} / {second['title']}",
                            "shared_terms": ", ".join(sorted(shared_terms)[:8]),
                            "paper_a_idea": first_idea,
                            "paper_b_idea": second_idea,
                        }
                    )
                    break
                if len(common_ideas) >= 8:
                    break

    seen_themes = set()
    common_themes = []
    for idea in common_ideas:
        if idea["theme"] in seen_themes:
            continue
        seen_themes.add(idea["theme"])
        common_themes.append(
            {
                "theme": idea["theme"],
                "shared_terms": idea["shared_terms"],
                "papers": idea["papers"],
            }
        )

    oppositions = []
    increase_terms = {"increase", "increased", "higher", "elevated", "upregulated"}
    decrease_terms = {"decrease", "decreased", "lower", "reduced", "downregulated"}
    for first_index, first in enumerate(prepared):
        first_results = " ".join(first["results"]).lower()
        for second in prepared[first_index + 1 :]:
            second_results = " ".join(second["results"]).lower()
            first_up = any(term in first_results for term in increase_terms)
            first_down = any(term in first_results for term in decrease_terms)
            second_up = any(term in second_results for term in increase_terms)
            second_down = any(term in second_results for term in decrease_terms)
            shared = sorted((first["keywords"] & second["keywords"]) or (set(shared_keywords) & all_keywords))
            if shared and ((first_up and second_down) or (first_down and second_up)):
                oppositions.append(
                    {
                        "papers": f"{first['title']} / {second['title']}",
                        "shared_terms": ", ".join(shared[:6]),
                        "note": "One paper contains increase-like result language while the other contains decrease-like result language.",
                    }
                )

    comparison = {
        "papers": prepared,
        "shared_keywords": shared_keywords,
        "common_themes": common_themes,
        "common_ideas": common_ideas,
        "unique_keywords": unique_keywords,
        "oppositions": oppositions,
        "comparison_provider": "Local extractor",
    }

    deepseek_comparison = call_deepseek_comparison(papers)
    if deepseek_comparison:
        comparison.update(deepseek_comparison)
    return comparison


def build_reading_assistant(paper_text: str) -> dict[str, Any]:
    """Create local reading support from a pasted or uploaded paper.

    This is intentionally simple and transparent. It extracts useful points; it
    does not claim to replace careful reading or expert judgement.
    """
    cleaned = clean_paper_text(paper_text)
    sections = extract_sections(cleaned)

    method_terms = ["sample", "participants", "cells", "mice", "assay", "measured", "analysis"]
    result_terms = [
        "increased",
        "decreased",
        "significant",
        "associated",
        "correlated",
        "higher",
        "lower",
        "reduced",
        "improved",
    ]
    limitation_terms = ["limitation", "limited", "bias", "small sample", "future", "not assess"]

    local_summary = {
        "cleaned_text": cleaned,
        "word_count": len(re.findall(r"\b\w+\b", cleaned)),
        "sections": sections,
        "sources": extract_sources(cleaned),
        "keywords": extract_keywords(cleaned),
        "short_summary": generate_reworded_summary(sections, cleaned),
        "key_points": generate_reworded_points(sections, cleaned),
        "method_points": find_sentences_with_terms(sections.get("methods") or cleaned, method_terms, limit=5),
        "result_points": find_sentences_with_terms(sections.get("results") or cleaned, result_terms, limit=5),
        "limitation_points": find_sentences_with_terms(cleaned, limitation_terms, limit=5),
        "summary_provider": "Local extractor",
    }

    deepseek_summary = call_deepseek_summary(cleaned)
    if not deepseek_summary:
        return local_summary

    for field in [
        "short_summary",
        "key_points",
        "keywords",
        "method_points",
        "result_points",
        "limitation_points",
        "summary_provider",
    ]:
        if deepseek_summary.get(field):
            local_summary[field] = deepseek_summary[field]
    return local_summary


def _line(label: str, value: Any) -> str:
    """Format one labeled summary line."""
    if isinstance(value, list):
        value = ", ".join(str(item) for item in value if str(item).strip())
    return f"{label}: {value or 'Not entered'}"


def generate_summary(paper: dict[str, Any], vocabulary: Optional[pd.DataFrame] = None) -> str:
    """Build a plain-language structured summary from the current paper review."""
    results = paper.get("results", [])
    result_lines = []
    for index, result in enumerate(results, start=1):
        if not any(str(value).strip() for value in result.values()):
            continue
        result_lines.append(
            "\n".join(
                [
                    f"Result {index}",
                    _line("Description", result.get("description")),
                    _line("Direction", result.get("direction")),
                    _line("Statistical significance", result.get("significance")),
                    _line("Figure/Table", result.get("figure")),
                    _line("Simple explanation", result.get("explanation")),
                ]
            )
        )

    methods_summary = "\n".join(
        [
            _line("Study type", paper.get("study_type")),
            _line("Sample used", paper.get("sample_used")),
            _line("Sample size", paper.get("sample_size")),
            _line("Variables measured", paper.get("variables_measured")),
            _line("Techniques used", paper.get("techniques_used", [])),
            _line("Controls used", paper.get("controls_used")),
            _line("Equipment used", paper.get("equipment_used")),
            _line("Statistical methods", paper.get("statistical_methods")),
            _line("Ethical approval mentioned", paper.get("ethical_approval")),
            _line("Method notes", paper.get("method_notes")),
        ]
    )

    vocabulary_lines = ["Not entered"]
    if vocabulary is not None and not vocabulary.empty:
        title = paper.get("paper_title", "")
        related_vocab = vocabulary[
            vocabulary["Source paper title"].astype(str).str.lower() == title.lower()
        ]
        if related_vocab.empty:
            related_vocab = vocabulary.head(10)
        vocabulary_lines = [
            f"- {row['Term']}: {row['Simple explanation']}"
            for _, row in related_vocab.iterrows()
            if str(row["Term"]).strip()
        ] or ["Not entered"]

    citation = ", ".join(
        part
        for part in [
            paper.get("authors"),
            f"({paper.get('year')})" if paper.get("year") else "",
            paper.get("paper_title"),
            paper.get("journal_name"),
            paper.get("doi_link"),
        ]
        if part
    )

    sections = [
        ("Research Paper Reading Helper Summary", ""),
        ("Paper title", paper.get("paper_title", "Not entered")),
        ("Full citation-style information", citation or "Not entered"),
        ("Extracted short summary", paper.get("auto_summary", "Not entered")),
        ("Extracted content points", "\n".join(f"- {point}" for point in paper.get("auto_key_points", [])) or "Not entered"),
        ("One-sentence summary", paper.get("one_sentence_summary", "Not entered")),
        ("Background", paper.get("background_information", "Not entered")),
        ("Central study focus", paper.get("research_question", "Not entered")),
        ("Aim", paper.get("study_aim", "Not entered")),
        ("Expected finding", paper.get("hypothesis", "Not entered")),
        ("Study type", paper.get("study_type", "Not entered")),
        ("Methods summary", methods_summary),
        ("Key results", "\n\n".join(result_lines) or "Not entered"),
        ("Discussion summary", paper.get("results_meaning", "Not entered")),
        ("Main conclusion", paper.get("main_conclusion", "Not entered")),
        (
            "Limitations",
            "\n".join(
                [
                    _line("Checklist", paper.get("common_limitations", [])),
                    _line("Author-stated limitations", paper.get("author_limitations")),
                    _line("Additional limitations", paper.get("student_limitations")),
                ]
            ),
        ),
        ("Important vocabulary", "\n".join(vocabulary_lines)),
        ("Revision notes", paper.get("questions_still_have", "Not entered")),
    ]

    return "\n\n".join(
        f"{heading}\n{'=' * len(heading)}\n{body}" if body else heading
        for heading, body in sections
    )


def _export_path(paper: dict[str, Any], extension: str) -> Path:
    """Create a standard export path for one paper and file extension."""
    ensure_storage_files()
    title = clean_filename(paper.get("paper_title", "untitled_paper"))
    reviewed = paper.get("date_reviewed") or date.today().isoformat()
    return EXPORTS_DIR / f"research_summary_{title}_{reviewed}.{extension}"


def export_summary_txt(summary: str, paper: dict[str, Any]) -> Path:
    """Export a summary as a text file."""
    path = _export_path(paper, "txt")
    path.write_text(summary, encoding="utf-8")
    return path


def export_summary_json(paper: dict[str, Any], summary: str) -> Path:
    """Export a paper review and its generated summary as JSON."""
    path = _export_path(paper, "json")
    payload = {"paper": paper, "summary": summary}
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    return path


def export_summary_csv(paper: dict[str, Any], summary: str) -> Path:
    """Export a one-row CSV containing the paper title and generated summary."""
    path = _export_path(paper, "csv")
    row = {
        "paper_id": paper.get("paper_id", ""),
        "paper_title": paper.get("paper_title", ""),
        "date_reviewed": paper.get("date_reviewed", ""),
        "summary": summary,
    }
    pd.DataFrame([row]).to_csv(path, index=False)
    return path


def export_summary_docx(summary: str, paper: dict[str, Any]) -> Path:
    """Export a summary as a Word document."""
    if Document is None:
        raise ImportError("python-docx is required for DOCX export.")
    path = _export_path(paper, "docx")
    document = Document()
    document.add_heading("Research Paper Reading Helper Summary", level=1)
    for block in summary.split("\n\n"):
        lines = block.splitlines()
        if len(lines) >= 2 and set(lines[1]) == {"="}:
            document.add_heading(lines[0], level=2)
            text = "\n".join(lines[2:]).strip()
            if text:
                document.add_paragraph(text)
        else:
            document.add_paragraph(block)
    document.save(path)
    return path
